"""DOCX parser — python-docx. Trích paragraphs + tables."""

from __future__ import annotations

from pathlib import Path

from app.pipeline.parsers import ParsedDoc


def parse(path: Path) -> ParsedDoc:
    from docx import Document  # lazy import

    doc = Document(str(path))
    paragraphs: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            style = (p.style.name or "").lower() if p.style else ""
            if style.startswith("heading 1"):
                paragraphs.append(f"# {p.text}")
            elif style.startswith("heading 2"):
                paragraphs.append(f"## {p.text}")
            elif style.startswith("heading 3"):
                paragraphs.append(f"### {p.text}")
            else:
                paragraphs.append(p.text)

    tables: list[list[list[str]]] = []
    for t in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in t.rows]
        if rows:
            tables.append(rows)

    return ParsedDoc(
        text="\n\n".join(paragraphs),
        tables=tables,
        metadata={"source_type": "docx"},
    )
